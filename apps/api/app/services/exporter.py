"""Markdown → md/docx/pdf rendering.

One module, three pure-Python renderers (no system deps):
- markdown bytes   — pass-through
- DOCX (python-docx) — paragraphs + headings, code blocks, bullet/numbered lists
- PDF  (reportlab)   — flowables built from the same markdown AST

Both DOCX and PDF use the `markdown` library to tokenize so they share heading /
list / code-block semantics with the rendered web view.
"""
from __future__ import annotations

import io
import re
from html.parser import HTMLParser
from typing import Literal

import markdown as md_lib
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
)

ExportFormat = Literal["md", "docx", "pdf"]

MIME = {
    "md": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def safe_filename(title: str, fmt: ExportFormat) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", title).strip("_") or "export"
    return f"{slug}.{fmt}"


# ---------------------------------------------------------------------------
# Tiny block-level walker: turn markdown text into an ordered list of blocks
# that both DOCX and PDF renderers consume. We sit on top of python-markdown's
# HTML output and pull out only the block types we render.
# ---------------------------------------------------------------------------

class _HTMLBlockExtractor(HTMLParser):
    """Walks markdown-rendered HTML and yields normalized block dicts."""

    BLOCK_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "pre", "ul", "ol", "blockquote"}

    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[dict] = []
        self._stack: list[str] = []
        self._buffer: list[str] = []
        self._list_stack: list[dict] = []  # tracks current ul/ol with items

    # ------- helpers
    def _flush_text(self) -> str:
        text = "".join(self._buffer).strip()
        self._buffer = []
        return text

    def _top(self) -> str | None:
        return self._stack[-1] if self._stack else None

    # ------- parser hooks
    def handle_starttag(self, tag, attrs):
        if tag in self.BLOCK_TAGS:
            self._stack.append(tag)
            if tag in ("ul", "ol"):
                self._list_stack.append({"type": tag, "items": []})
            self._buffer = []
        elif tag == "li":
            self._stack.append(tag)
            self._buffer = []
        elif tag in ("strong", "b", "em", "i", "code"):
            # inline — render as plain text for now (keeps PDF/DOCX simple)
            pass

    def handle_endtag(self, tag):
        if tag == "li":
            text = self._flush_text()
            if self._list_stack:
                self._list_stack[-1]["items"].append(text)
            if self._stack and self._stack[-1] == "li":
                self._stack.pop()
            return

        if tag in ("ul", "ol"):
            lst = self._list_stack.pop() if self._list_stack else None
            if lst:
                self.blocks.append({"kind": "list", "ordered": tag == "ol", "items": lst["items"]})
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()
            return

        if tag in self.BLOCK_TAGS:
            text = self._flush_text()
            if tag.startswith("h"):
                self.blocks.append({"kind": "heading", "level": int(tag[1]), "text": text})
            elif tag == "p" and text:
                self.blocks.append({"kind": "paragraph", "text": text})
            elif tag == "pre" and text:
                self.blocks.append({"kind": "code", "text": text})
            elif tag == "blockquote" and text:
                self.blocks.append({"kind": "quote", "text": text})
            if self._stack and self._stack[-1] == tag:
                self._stack.pop()

    def handle_data(self, data):
        self._buffer.append(data)


def _markdown_to_blocks(text: str) -> list[dict]:
    html = md_lib.markdown(text, extensions=["fenced_code", "tables"])
    extractor = _HTMLBlockExtractor()
    extractor.feed(html)
    return extractor.blocks


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def render_docx(*, title: str, markdown_text: str) -> bytes:
    doc = Document()

    # Body font defaults — Calibri 11pt is Word's built-in default
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    for block in _markdown_to_blocks(markdown_text):
        kind = block["kind"]
        if kind == "heading":
            level = min(max(block["level"], 1), 4)
            doc.add_heading(block["text"], level=level)
        elif kind == "paragraph":
            doc.add_paragraph(block["text"])
        elif kind == "quote":
            p = doc.add_paragraph(block["text"])
            p.paragraph_format.left_indent = Pt(18)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x5C, 0x58, 0x49)
        elif kind == "list":
            style_name = "List Number" if block["ordered"] else "List Bullet"
            for item in block["items"]:
                doc.add_paragraph(item, style=style_name)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------

def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=base["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        textColor="#1F1E1A",
        spaceAfter=8,
        alignment=TA_LEFT,
    )
    return {
        "title": ParagraphStyle("Title", parent=base["Title"], fontName="Helvetica-Bold",
                                 fontSize=22, leading=26, textColor="#1F1E1A", spaceAfter=14),
        "h1": ParagraphStyle("H1", parent=base["Heading1"], fontName="Helvetica-Bold",
                              fontSize=18, leading=22, textColor="#1F1E1A", spaceBefore=12, spaceAfter=8),
        "h2": ParagraphStyle("H2", parent=base["Heading2"], fontName="Helvetica-Bold",
                              fontSize=15, leading=19, textColor="#1F1E1A", spaceBefore=10, spaceAfter=6),
        "h3": ParagraphStyle("H3", parent=base["Heading3"], fontName="Helvetica-Bold",
                              fontSize=13, leading=17, textColor="#3D3929", spaceBefore=8, spaceAfter=4),
        "body": body,
        "quote": ParagraphStyle("Quote", parent=body, leftIndent=18, fontName="Helvetica-Oblique",
                                 textColor="#5C5849", borderColor="#C96442", borderPadding=0),
        "code": ParagraphStyle("Code", parent=base["Code"], fontName="Courier", fontSize=9,
                                leading=12, leftIndent=12, backColor="#F5F3EE", borderPadding=6,
                                spaceBefore=4, spaceAfter=8),
    }


def _escape_pdf(text: str) -> str:
    # reportlab Paragraph uses XML-ish markup — escape bare ampersands and angle brackets.
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_pdf(*, title: str, markdown_text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        title=title,
    )
    styles = _pdf_styles()
    flow: list = [Paragraph(_escape_pdf(title), styles["title"]), Spacer(1, 0.1 * inch)]

    for block in _markdown_to_blocks(markdown_text):
        kind = block["kind"]
        if kind == "heading":
            level = max(1, min(block["level"], 3))
            flow.append(Paragraph(_escape_pdf(block["text"]), styles[f"h{level}"]))
        elif kind == "paragraph":
            flow.append(Paragraph(_escape_pdf(block["text"]), styles["body"]))
        elif kind == "quote":
            flow.append(Paragraph(_escape_pdf(block["text"]), styles["quote"]))
        elif kind == "list":
            items = [ListItem(Paragraph(_escape_pdf(t), styles["body"]), leftIndent=12)
                     for t in block["items"]]
            flow.append(ListFlowable(
                items,
                bulletType="1" if block["ordered"] else "bullet",
                start="1" if block["ordered"] else None,
                bulletFontName="Helvetica",
                bulletFontSize=11,
                leftIndent=18,
            ))
            flow.append(Spacer(1, 0.05 * inch))
        elif kind == "code":
            flow.append(Preformatted(block["text"], styles["code"]))

    doc.build(flow)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public render() dispatcher
# ---------------------------------------------------------------------------

def render(*, fmt: ExportFormat, title: str, markdown_text: str) -> bytes:
    if fmt == "md":
        return markdown_text.encode("utf-8")
    if fmt == "docx":
        return render_docx(title=title, markdown_text=markdown_text)
    if fmt == "pdf":
        return render_pdf(title=title, markdown_text=markdown_text)
    raise ValueError(f"Unsupported export format: {fmt}")


# ---------------------------------------------------------------------------
# Conversation transcript → markdown source
# ---------------------------------------------------------------------------

def conversation_to_markdown(*, title: str, messages: list) -> str:
    """Render a conversation as markdown. `messages` is a list of ORM Message rows
    or anything with .role and .content attributes."""
    lines = [f"# {title}", ""]
    for m in messages:
        role = (getattr(m, "role", "") or "").upper() or "MESSAGE"
        content = (getattr(m, "content", "") or "").rstrip()
        lines.append(f"## {role}")
        lines.append("")
        lines.append(content)
        lines.append("")
    return "\n".join(lines)
